import importlib.util
import os
import tempfile
import unittest
from pathlib import Path

from shadow_text.engine import Span
from shadow_text.watcher import process_for_censura, process_for_riunione


class EmailDetector:
    def detect(self, text: str) -> list[Span]:
        value = "test@example.com"
        start = text.index(value)
        return [Span(label="private_email", start=start, end=start + len(value), text=value)]


@unittest.skipUnless(importlib.util.find_spec("docx"), "python-docx non installato")
class DocxWorkflowTests(unittest.TestCase):
    def test_docx_censor_and_restore_keep_docx_format(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            censura = root / "Censura"
            dati = root / "Dati"
            riunione = root / "Riunione"
            censura.mkdir()
            dati.mkdir()
            riunione.mkdir()
            source = censura / "nota.docx"
            document = Document()
            document.add_paragraph("Email: test@example.com")
            document.save(str(source))

            with _temporary_key("azienda-secret"):
                redacted = process_for_censura(source, data_dir=dati, detector=EmailDetector())
                assert redacted is not None
                self.assertEqual(redacted.suffix, ".docx")
                self.assertIn("EMAIL_00001", _docx_text(redacted))
                self.assertNotIn("test@example.com", _docx_text(redacted))

                riunione_file = riunione / redacted.name
                riunione_file.write_bytes(redacted.read_bytes())
                restored = process_for_riunione(riunione_file, data_dir=dati)

            assert restored is not None
            self.assertEqual(restored.suffix, ".docx")
            self.assertIn("test@example.com", _docx_text(restored))


@unittest.skipUnless(importlib.util.find_spec("openpyxl"), "openpyxl non installato")
class XlsxWorkflowTests(unittest.TestCase):
    def test_xlsx_censor_and_restore_keep_xlsx_format(self):
        from openpyxl import Workbook, load_workbook

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            censura = root / "Censura"
            dati = root / "Dati"
            riunione = root / "Riunione"
            censura.mkdir()
            dati.mkdir()
            riunione.mkdir()
            source = censura / "nota.xlsx"
            workbook = Workbook()
            workbook.active["A1"] = "Email: test@example.com"
            workbook.save(source)
            workbook.close()

            with _temporary_key("azienda-secret"):
                redacted = process_for_censura(source, data_dir=dati, detector=EmailDetector())
                assert redacted is not None
                self.assertEqual(redacted.suffix, ".xlsx")
                self.assertEqual(_xlsx_cell(redacted), "Email: EMAIL_00001")

                riunione_file = riunione / redacted.name
                riunione_file.write_bytes(redacted.read_bytes())
                restored = process_for_riunione(riunione_file, data_dir=dati)

            assert restored is not None
            self.assertEqual(restored.suffix, ".xlsx")
            self.assertEqual(_xlsx_cell(restored), "Email: test@example.com")

        load_workbook


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _xlsx_cell(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path)
    value = workbook.active["A1"].value
    workbook.close()
    return value


class _temporary_key:
    def __init__(self, value: str | None) -> None:
        self.value = value
        self.original = os.environ.get("SHADOW_TEXT_KEY")

    def __enter__(self):
        if self.value is None:
            os.environ.pop("SHADOW_TEXT_KEY", None)
        else:
            os.environ["SHADOW_TEXT_KEY"] = self.value

    def __exit__(self, exc_type, exc, tb):
        if self.original is None:
            os.environ.pop("SHADOW_TEXT_KEY", None)
        else:
            os.environ["SHADOW_TEXT_KEY"] = self.original


if __name__ == "__main__":
    unittest.main()
